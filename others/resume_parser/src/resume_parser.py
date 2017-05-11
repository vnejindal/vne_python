"""
Parses the resume document present in .docx file 

Prerequisites: 
pip install pypiwin32
pip install python-docx
"""
import os
import shutil
import re
from docx import Document
import win32com.client as win32

### Experiment Area ###
#print dir(document)
#print document.paragraphs
#print document.tables 
#print dir(document.sections)
#print dir(document.sections.count)
#print document.element
#print document.part
#document.save('test.docx')

    
### Experiment Area ###

# List of output files that got converted successfully
src_file_list = []
dst_file_list = []

def get_years_exp(line_str):
    """
    second level of parsing to get specific years of experience
    """
    """
    get the sentence from paragraph 
    break that sentence in list of words
    try to get number of years from that list
    
    """
    #print line_str
    
    keywords = ['years', 'experience']
    sentences = line_str.split('.')
    found = False
    
    wlist = []
    for sentence in sentences:
        for keyword in keywords:
            if sentence.find(keyword) is not -1:
                found = True
            else:
                found = False
        if found is True:
            #Sentence is found
            wlist = sentence.split()
            break
    if len(wlist) != 0:
        for word in wlist:
            #take care of scenarios like 8+ years of experience
            if word.find('+') != -1:
                word = word[:-1]
            if word.isdigit() is True:
                return word
    
    return None
    
def parse_para(line_str):
    """
    parses paragraph string to identify literals
    """
    keywords = ['years', 'experience']
    years = None
    
    if len(line_str) == 0:
        return years
    
    found = False
    for keyword in keywords:
        if line_str.lower().find(keyword) is not -1:
            found = True
        else:
            break;
    if found is True:
#        print line_str
        years = get_years_exp(line_str)
    
    return years
        

def read_resume(filename, extension):
    """
    reads resume to get information
    filename: Full path of file 
    extension: file extension  
    """
    years = None
    full_name = ''
    
    if extension == '.docx':
        document = Document(filename) 
       
        """
        retval = False
        count = 0
        while count is not len(document.paragraphs):
            #print count, document.paragraphs[count].text
            retval = parse_para(document.paragraphs[count].text)
            if retval is True:
                break
            count +=1
        """
        name_found = False
        for para in document.paragraphs:
            ## Find full name
            if len(para.text) != 0 and name_found is False:
                full_name = '_'.join(para.text.split())
                #print full_name
                name_found = True
                continue
                
            years = parse_para(para.text)
            if years is not None:
                break   

    elif extension == '.doc':
        """
        Convert .doc to .txt
        read .txt line by line
        """
        name_found = False
        ofile_name = convert_doc_file(filename)
        if len(ofile_name) != 0:
            file_fd = open(ofile_name)
            for line in file_fd.readlines():
                ## Find full name
                if len(line) != 0 and name_found is False:
                    full_name = '_'.join(line.split())
                    #print full_name
                    name_found = True
                    continue
                
                years = parse_para(line)
                if years is not None:
                    break   
                
            file_fd.close()
        else:
            #vne::tbd::
            pass
        
    retlist = []
    if years is not None:
        #print 'Experience Found', years, full_name
        retlist.append(full_name)
        retlist.append(years)
    else:
        print 'Not able to find experience info'
    
    #returns [full_name, years] or []
    return retlist

def output_resume(dirpath, filename, retlist):
    """
    
    """
    #create directory if it does not exist
    output_dir = 'output\\' + dirpath
    #vne::tbd:: It can be moved globally for performance optimization
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    extension = os.path.splitext(filename)[1]
    fname = '_'.join(retlist + [extension])
    
    srcfile = '\\'.join([dirpath, filename])
    dstfile = '\\'.join([output_dir, fname])
    #Remove unwanted characters from filename
    dstfile= re.sub('[\-\\|/:(){}<>]', '', dstfile)
    print 'vne::', srcfile, dstfile
    
    global src_file_list
    global dst_file_list
    src_file_list.append(srcfile)
    dst_file_list.append(dstfile)
    
    shutil.copy(srcfile, dstfile)
    #shutil.move(srcfile, dstfile)
    
def browse_files(base_dir):
    """
    browses directory to identify files
    """
    if not os.path.exists(base_dir):
        print 'base directory doest not exist:', base_dir
        return
    
    retlist = []
    cur_dir = os.getcwd().split('\\')
    
    for (dirpath, dirnames, filenames) in os.walk(base_dir):
        #print 'processing directory: ', dirpath
        if len(filenames) != 0:
            #print filenames
            for filename in filenames:
                extension = os.path.splitext(filename)[1]
                if extension.lower() == '.docx' or extension.lower() == '.doc':
                    fname = '\\\\'.join(cur_dir + [dirpath, filename])
                    #print 'processing file: ', fname
                    try:
                    #if True:
                        retlist = read_resume(fname, extension.lower())
                        if len(retlist) != 0: 
                            #get designation 
                            print 'processing file: ', fname
                            #Get designation from file name
                            designation = dirpath.split('\\')[-1]
                            designation = '_'.join(designation.split())
                            #print 'desig: ', designation
                            output_resume(dirpath, filename, [designation] + retlist)
                    except:
                        print 'Exception Occured: ', fname



def convert_doc_file(file_name):
    """
    This function converts .doc file to its .txt format
    returns outfile .txt file name with absolute path
    """
    
    ofile_name=file_name.split('.')[0] + '.txt'
    #print 'processing:', file_name, ofile_name
    
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Documents.Open(file_name)
    try:
        word.ActiveDocument.SaveAs(ofile_name, FileFormat=win32.constants.wdFormatTextLineBreaks)
    except: 
        ofile_name = ''
        print 'could not convert .doc to .txt', file_name
    
    word.Documents.Close()
    return ofile_name

base_dir = 'resumes'
browse_files(base_dir)

print 'src files:'
print src_file_list
print 'dst files:'
print dst_file_list
print 'Total Files processed: ', len(src_file_list), len(dst_file_list)

#read_resume('resumes\\Artificial Intelligence\\AnandAndrew_J2EEJava_LeadResume.docx')
#read_resume('resume_goldsmith.docx')
#print os.walk('resumes')